#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build the WSQ assessment set for 'Application Integration with Docker and Kubernetes' (TGS-2021010366):
  - Written Assessment (SAQ)  — 5 open-ended KNOWLEDGE questions (K1–K5), aligned to the slides
  - Practical Performance (PP) — 4 PRACTICAL tasks (LO1–LO4), aligned to the in-class activities
Each instrument is produced as a Question Paper and a matching Answer Key (4 DOCX total),
all with the WSQ house cover page (same as the Lesson Plan / Learner Guide). Page 1 is the cover;
page 2 carries Trainee Information + Instructions + Grading; the questions/tasks begin on page 3.
Body: Arial 11.
"""
import os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# This script lives in the wsq-assessment skill (.claude/skills/wsq-assessment/) and runs in
# place — it detects the course repo root by walking up to the nearest dir that has a .git
# folder (or both courseware/ and assessment/). Override with env REPO=/path if needed.
def _find_repo():
    env = os.environ.get("REPO")
    if env and os.path.isdir(env):
        return os.path.abspath(env)
    d = os.path.dirname(os.path.abspath(__file__))
    while d != os.path.dirname(d):
        if os.path.isdir(os.path.join(d, ".git")) or \
           (os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "assessment"))):
            return d
        d = os.path.dirname(d)
    return os.getcwd()

REPO = _find_repo()
# prodoc.py (WSQ cover page + version control + page numbers, same as LP/LG) ships with the
# tertiary-lesson-plan skill. Look for it at the project level first, then the user level.
for _cand in (os.path.join(REPO, ".claude/skills/tertiary-lesson-plan"),
              os.path.expanduser("~/.claude/skills/tertiary-lesson-plan")):
    if os.path.exists(os.path.join(_cand, "prodoc.py")):
        sys.path.insert(0, _cand); break
import prodoc  # cover page + version control + page numbers (same as LP/LG)

# ─── EDIT PER COURSE ────────────────────────────────────────────────────────
TITLE       = "Application Integration with Docker and Kubernetes"
COURSE_CODE = "TGS-2021010366"
# ────────────────────────────────────────────────────────────────────────────
# The cover page renders prodoc's module-level TGS constant. Override it so the
# assessment cover shows THIS course's ref (works with either prodoc version —
# the older project prodoc has no course_code kwarg).
prodoc.TGS = f"TGS Ref No: {COURSE_CODE}"
OUT   = os.path.join(REPO, "assessment")

# Logos: prefer the course's own courseware/assets, else fall back to the copies bundled
# in this skill (so the assessment builds even outside this project). Replace the course
# logo per course; the Tertiary Infotech logo is the same for every WSQ course.
def _logo(name):
    here = os.path.dirname(os.path.abspath(__file__))
    for p in (os.path.join(REPO, "courseware/assets", name), os.path.join(here, "assets", name)):
        if os.path.exists(p):
            return p
    return None
ORG_LOGO    = _logo("tertiary-infotech-logo.png")
COURSE_LOGO = _logo("docker-k8s-course-logo.png")   # None if absent → Tertiary-only cover (as LP/LG)

Q_VER, A_VER = "v5", "v5"   # single standardised version across all four files
BRAND = RGBColor(0x1F, 0x6F, 0xEB); DARK = RGBColor(0x11, 0x18, 0x27); GREY = RGBColor(0x55, 0x5B, 0x66)
# Assessments carry the cover page only — no Document Version Control Record.

# ---------------------------------------------------------------- WRITTEN (KNOWLEDGE)
# (criterion, context, question, [model-answer points]) — each traces to the course slides.
WRITTEN = [
 ("K1",
  "Containers have largely replaced virtual machines as the unit of application packaging. A VM virtualises "
  "hardware and runs a full guest operating system on a hypervisor, while a container packages just the "
  "application and its dependencies.",
  "What are the differences between a Virtual Machine and a container, and why do these differences matter in "
  "modern application development?",
  ["A VM virtualises hardware and runs a full guest OS on a hypervisor; a container shares the host OS kernel "
   "and packages only the application plus its dependencies.",
   "VMs are heavier: gigabytes in size, boot in seconds–minutes, and carry hypervisor overhead "
   "(e.g. VMware, VirtualBox, EC2).",
   "Containers are lighter: megabytes in size, start in milliseconds, with minimal overhead "
   "(e.g. Docker, containerd).",
   "Why it matters: containers pack many more workloads per host, give consistent 'build once, run anywhere' "
   "environments, and start/scale fast — ideal for microservices and CI/CD. "
   "(Slides: Why Containers — Virtual Machines vs Containers / How It Works — Docker Architecture)"]),
 ("K2",
  "A Docker image is built automatically from a text recipe. Getting the instructions and their order right "
  "controls both what the image contains and how fast it rebuilds.",
  "How do you create a Docker image, and what are the key instructions in a Dockerfile?",
  ["A Dockerfile is a text recipe of instructions Docker runs to build an image automatically; you build it "
   "with `docker build -t <name> .`.",
   "Key instructions: FROM (base image), WORKDIR, COPY, RUN, ENV, EXPOSE, VOLUME, and CMD / ENTRYPOINT.",
   "Each instruction creates a cached layer; unchanged layers are reused for fast rebuilds. Order matters — put "
   "rarely-changing steps first (COPY requirements.txt + RUN pip install BEFORE COPY . .) to keep dependencies "
   "cached, and use a .dockerignore to keep junk out of the build context.",
   "CMD sets a default, overridable command; ENTRYPOINT sets a fixed executable with args appended. "
   "(Slides: Building Images — What is a Dockerfile? / Layers & Cache / CMD vs ENTRYPOINT)"]),
 ("K3",
  "Docker runs containers on a single host. As applications grow you need to run and manage many containers "
  "reliably across a cluster of machines.",
  "What are the benefits of using Kubernetes for deploying and managing containerised applications?",
  ["Docker runs containers on one host; Kubernetes orchestrates them across a cluster of nodes.",
   "It self-heals (restarts failed Pods), scales workloads up and down, and performs rolling updates with no "
   "downtime.",
   "It is declarative: you describe the desired state in YAML and Kubernetes continuously makes it true; "
   "kubectl is the CLI (get / describe / apply / delete / scale / rollout).",
   "Architecture: a Control Plane (API Server, etcd, Scheduler, Controller Manager) drives Worker Nodes "
   "(kubelet, kube-proxy) that run your Pods. "
   "(Slides: Orchestration — Why Kubernetes? / How It Works — Cluster Architecture)"]),
 ("K4",
  "In Kubernetes a Pod's IP address changes every time it restarts, scales or is updated, so clients cannot "
  "target Pods directly. Services solve this.",
  "Why do Pods need Services, and what are the main types of Service in Kubernetes and how does each function?",
  ["Pod IPs change on every restart, scale or update, so you can't rely on them; a Service gives a stable "
   "address and load-balances across the Pods it selects (by labels).",
   "ClusterIP — internal only; reachable from inside the cluster for service-to-service traffic (the default).",
   "NodePort — opens a fixed port (30000–32767) on every node, giving external access to the Service.",
   "LoadBalancer — provisions an external (cloud) load balancer with a single external IP, extending NodePort; "
   "the labs use ClusterIP and NodePort. "
   "(Slide: Networking — Why Services?)"]),
 ("K5",
  "Not every workload runs forever. Alongside long-running Deployments, Kubernetes provides objects for work "
  "that finishes.",
  "What is the difference between a Job and a CronJob in Kubernetes, and when should each be used?",
  ["A Job runs one or more Pods to completion — it runs to success and then stops, and does not restart on "
   "success; use it for finite batch work (e.g. a one-off TaskBoard report or a database migration).",
   "A CronJob runs Jobs on a schedule using standard cron syntax; use it for recurring scheduled work "
   "(e.g. a nightly cleanup or backup).",
   "Contrast with a Deployment, which keeps long-running Pods alive indefinitely (like a web server) — "
   "Jobs and CronJobs are for work that is meant to finish. "
   "(Slides: Persist & Schedule — Storage & Batch Workloads / Lab 19 — Jobs & CronJobs)"]),
]

# ---------------------------------------------------------------- PRACTICAL (ACTIVITY-BASED)
SCENARIO = (
 "You have joined a startup as a DevOps engineer. The team is shipping a small web application and you own the "
 "full path from a single container to a Kubernetes cluster: containerise a Flask app and publish it to Docker "
 "Hub, stand up a multi-service stack with Docker Compose, then migrate the workload to Kubernetes — deploying "
 "and troubleshooting a Pod, exposing it with a Service, and giving it persistent storage. Complete the four "
 "tasks below; each mirrors a hands-on activity you did in class. For each task, paste your Dockerfile / YAML "
 "and a screenshot of your output as evidence.")

# (label, criterion, task prompt, box caption, model-answer build steps citing the activity)
BOX_CAP = "Paste your Dockerfile / YAML and a screenshot of your output in the box below"
PRACTICAL = [
 ("Task 1", "LO1",
  "Containerise a Flask application (Docker). You are given a simple Flask note-taking app (app.py and "
  "requirements.txt). "
  "Part A — Write a Dockerfile that: uses python:3.12-slim; sets WORKDIR to /app; copies requirements.txt and "
  "installs deps with pip install --no-cache-dir; copies the rest of the code; sets ENV DATA_DIR=/app/data and "
  "APP_PORT=5000; declares /app/data as a VOLUME; EXPOSEs 5000; and runs the app with python app.py. "
  "Part B — Build the image as notes-app, run it with the container port mapped to 5001 on the host, add a note "
  "with curl -X POST -d \"note=Hello Docker\" http://localhost:5001/add and view it with curl "
  "http://localhost:5001/notes. "
  "Part C — Tag the image as <your-username>/notes-app:v1, push it to Docker Hub, and state the command someone "
  "else would run to pull and use it. (Labs 3-4 — Build image · Lab 9 — Docker Hub.)",
  BOX_CAP,
  "Part A — Dockerfile:\n"
  "FROM python:3.12-slim\n"
  "WORKDIR /app\n"
  "COPY requirements.txt .\n"
  "RUN pip install --no-cache-dir -r requirements.txt\n"
  "COPY . .\n"
  "ENV DATA_DIR=/app/data\n"
  "ENV APP_PORT=5000\n"
  "VOLUME /app/data\n"
  "EXPOSE 5000\n"
  "CMD [\"python\", \"app.py\"]\n\n"
  "Part B — Build & run:\n"
  "docker build -t notes-app .\n"
  "docker run -d -p 5001:5000 notes-app\n"
  "curl -X POST -d \"note=Hello Docker\" http://localhost:5001/add\n"
  "curl http://localhost:5001/notes\n\n"
  "Part C — Push to Docker Hub:\n"
  "docker tag notes-app <your-username>/notes-app:v1\n"
  "docker login\n"
  "docker push <your-username>/notes-app:v1\n"
  "# Someone else pulls & runs:\n"
  "docker pull <your-username>/notes-app:v1\n"
  "docker run -p 5001:5000 <your-username>/notes-app:v1"),
 ("Task 2", "LO2",
  "Deploy a multi-service site with Docker Compose. Stand up a WordPress site backed by a MySQL database. "
  "Part A — Write a docker-compose.yml with two services: db using image mysql:8.0 with env "
  "MYSQL_ROOT_PASSWORD=rootpass, MYSQL_DATABASE=wordpress, MYSQL_USER=wpuser, MYSQL_PASSWORD=wppass and a named "
  "volume db-data mounted at /var/lib/mysql; and wordpress using image wordpress:latest, mapping host port 8080 "
  "to container port 80, with env WORDPRESS_DB_HOST=db, WORDPRESS_DB_USER=wpuser, WORDPRESS_DB_PASSWORD=wppass, "
  "WORDPRESS_DB_NAME=wordpress, and depends_on db. "
  "Part B — Start the stack with docker compose up -d, confirm both services with docker compose ps, and open "
  "http://localhost:8080 to reach the WordPress setup page. (Labs 10-12 — Docker Compose.)",
  BOX_CAP,
  "Part A — docker-compose.yml:\n"
  "services:\n"
  "  db:\n"
  "    image: mysql:8.0\n"
  "    environment:\n"
  "      MYSQL_ROOT_PASSWORD: rootpass\n"
  "      MYSQL_DATABASE: wordpress\n"
  "      MYSQL_USER: wpuser\n"
  "      MYSQL_PASSWORD: wppass\n"
  "    volumes:\n"
  "      - db-data:/var/lib/mysql\n"
  "  wordpress:\n"
  "    image: wordpress:latest\n"
  "    ports:\n"
  "      - \"8080:80\"\n"
  "    environment:\n"
  "      WORDPRESS_DB_HOST: db\n"
  "      WORDPRESS_DB_USER: wpuser\n"
  "      WORDPRESS_DB_PASSWORD: wppass\n"
  "      WORDPRESS_DB_NAME: wordpress\n"
  "    depends_on:\n"
  "      - db\n"
  "volumes:\n"
  "  db-data:\n\n"
  "Part B — Run & verify:\n"
  "docker compose up -d\n"
  "docker compose ps        # both services Up\n"
  "# Browse http://localhost:8080 -> WordPress setup page"),
 ("Task 3", "LO3",
  "Deploy and troubleshoot a Pod in Kubernetes. "
  "1) Create the namespace ckad-prep. 2) In it, create a Pod named mypod with image nginx:2.3.5 exposing port "
  "80. 3) Identify why the container will not start and write the root cause to pod-error.txt. 4) Change the "
  "Pod's image to nginx:1.15.12. 5) List the Pod and confirm it is Running. 6) Shell into the container, run ls, "
  "note the output, and exit. 7) Retrieve the Pod's IP address. 8) Run a temporary busybox Pod, shell in and "
  "wget the nginx Pod on port 80. 9) Show the logs of mypod. 10) Delete the Pod and the namespace. "
  "(Labs 13-14 — Pods & Namespaces.)",
  BOX_CAP,
  "1. Create namespace:\n"
  "kubectl create namespace ckad-prep\n"
  "2. Create the Pod (bad image):\n"
  "kubectl run mypod --image=nginx:2.3.5 --port=80 -n ckad-prep\n"
  "3. Diagnose — image nginx:2.3.5 does not exist -> ImagePullBackOff:\n"
  "kubectl get pod -n ckad-prep            # STATUS: ImagePullBackOff\n"
  "kubectl describe pod mypod -n ckad-prep # Events: manifest not found\n"
  "echo \"Image nginx:2.3.5 does not exist on Docker Hub.\" > pod-error.txt\n"
  "4. Fix the image:\n"
  "kubectl set image pod mypod mypod=nginx:1.15.12 -n ckad-prep\n"
  "5. Verify Running:\n"
  "kubectl get pod -n ckad-prep            # STATUS: Running\n"
  "6. Shell in (run ls, then exit):\n"
  "kubectl exec -it mypod -n ckad-prep -- /bin/sh\n"
  "7. Pod IP:\n"
  "kubectl get pods -o wide -n ckad-prep\n"
  "8. wget from a temporary busybox Pod:\n"
  "kubectl run busybox --image=busybox --rm -it --restart=Never -n ckad-prep -- wget -O- <pod-ip>:80\n"
  "9. Render logs:\n"
  "kubectl logs mypod -n ckad-prep\n"
  "10. Clean up:\n"
  "kubectl delete pod mypod -n ckad-prep\n"
  "kubectl delete namespace ckad-prep"),
 ("Task 4", "LO4",
  "Kubernetes Services and persistent storage. "
  "Part A — Create a deployment named myapp with 2 replicas of image nginx exposing container port 80. Expose "
  "it so it is reachable from inside the cluster, and verify with a temporary busybox Pod running wget against "
  "the Service. Then change the Service type so the Pods are reachable from outside the cluster and wget it "
  "from outside. "
  "Part B — Create a PersistentVolume my-pv of 1Gi using hostPath /tmp/k8s-data, and a PersistentVolumeClaim "
  "my-pvc requesting 500Mi; verify the PVC is Bound. Create a Pod storage-pod (image busybox) that mounts the "
  "PVC at /data and writes \"hello from storage\" to /data/message.txt. Delete and recreate the Pod, verify the "
  "data persists, then clean up all resources. (Labs 15 — Deployments · 17 — Services · 18 — Storage.)",
  BOX_CAP,
  "Part A — Routing traffic:\n"
  "kubectl create deployment myapp --image=nginx --replicas=2 --port=80\n"
  "kubectl expose deployment myapp --port=80 --target-port=80          # ClusterIP\n"
  "kubectl run tmp --image=busybox --rm -it --restart=Never -- wget -O- myapp:80\n"
  "# External access — switch to NodePort:\n"
  "kubectl delete service myapp\n"
  "kubectl expose deployment myapp --type=NodePort --port=80 --target-port=80\n"
  "kubectl get svc myapp                      # note the 3xxxx NodePort\n"
  "wget -O- localhost:<NodePort>\n\n"
  "Part B — Persistent storage:\n"
  "# pv.yaml\n"
  "apiVersion: v1\n"
  "kind: PersistentVolume\n"
  "metadata: { name: my-pv }\n"
  "spec:\n"
  "  capacity: { storage: 1Gi }\n"
  "  accessModes: [ReadWriteOnce]\n"
  "  hostPath: { path: /tmp/k8s-data }\n"
  "# pvc.yaml\n"
  "apiVersion: v1\n"
  "kind: PersistentVolumeClaim\n"
  "metadata: { name: my-pvc }\n"
  "spec:\n"
  "  accessModes: [ReadWriteOnce]\n"
  "  resources: { requests: { storage: 500Mi } }\n"
  "kubectl apply -f pv.yaml && kubectl apply -f pvc.yaml\n"
  "kubectl get pv,pvc                          # STATUS: Bound\n"
  "# storage-pod mounts my-pvc at /data and writes message.txt\n"
  "kubectl apply -f pod.yaml\n"
  "kubectl exec storage-pod -- cat /data/message.txt   # hello from storage\n"
  "kubectl delete pod storage-pod && kubectl apply -f pod.yaml\n"
  "kubectl exec storage-pod -- cat /data/message.txt   # data persists\n"
  "kubectl delete pod storage-pod; kubectl delete pvc my-pvc; kubectl delete pv my-pv"),
]

# ---------------------------------------------------------------- doc helpers
def base_doc():
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Arial"; n.font.size = Pt(11)
    return doc

def para(doc, text, size=11, bold=False, italic=False, color=None, after=6, before=0, align=None):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    if align is not None: p.alignment = align
    return p

def heading(doc, text, size=13):
    para(doc, text, size=size, bold=True, color=BRAND, after=6, before=8)

def answer_box(doc, lines=None, code=None, height_pt=90):
    """1x1 bordered box. `lines` → bullet-style model answer; `code` → monospace
    code/YAML/command block (indentation preserved); neither → empty answer space."""
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.paragraphs[0].text = ""
    if code:
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True; run.font.size = Pt(10.5)
        for ln in code.split("\n"):
            b = cell.add_paragraph(style=None)
            b.paragraph_format.space_after = Pt(0); b.paragraph_format.space_before = Pt(0)
            rr = b.add_run(ln if ln else " ")
            rr.font.name = "Consolas"; rr.font.size = Pt(9)
            rr._element.rPr.rFonts.set(qn('w:cs'), "Consolas")
            wt = rr._element.find(qn('w:t'))
            if wt is not None: wt.set(qn('xml:space'), 'preserve')
    elif lines:
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True; run.font.size = Pt(10.5)
        for ln in lines:
            b = cell.add_paragraph(style=None); b.paragraph_format.left_indent = Inches(0.15)
            rr = b.add_run("•  " + ln); rr.font.size = Pt(10.5)
    else:
        # empty answer space
        tr = t.rows[0]._tr
        trPr = tr.get_or_add_trPr(); trh = OxmlElement('w:trHeight')
        trh.set(qn('w:val'), str(int(height_pt*20))); trh.set(qn('w:hRule'), 'atLeast'); trPr.append(trh)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

FILL_GAP = 6    # extra space below each fill-in line (paired with double line spacing for writing room)

def candidate_block(doc):
    heading(doc, "Trainee Information")
    for label in ["Trainee Name (as per NRIC): ______________________________________",
                  "Last 3 digits and alphabet of NRIC/FIN: ____________________",
                  "Date: ____________________"]:
        p = para(doc, label, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

# Assessment briefing (from the course slides — "Briefing for Assessment").
BRIEFING = [
    "Place phones and other materials under the table or on the floor.",
    "No photos or recording of assessment scripts.",
    "No discussion during the assessment.",
    "Use a black/blue pen for hard-copy assessments.",
    "No liquid paper / correction tape.",
    "Scripts are collected when time is up.",
]

LMS_URL = "https://lms-tms.tertiaryinfotech.com/"

def add_hyperlink(p, url, text):
    """Add a real clickable Word hyperlink (blue, underlined) to paragraph p."""
    r_id = p.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink"); link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "22"); rPr.append(sz)  # 11pt
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    link.append(run); p._p.append(link)
    return link

def instructions(doc, minutes_text):
    heading(doc, "Instructions to Candidate")
    # None marks the upload instruction, which carries a clickable LMS hyperlink.
    items = [
        "This is an individual exercise.",
        "This is an open-book assessment.",
        f"A total of {minutes_text} is given to complete this assessment.",
        None,
    ] + BRIEFING
    for i, s in enumerate(items, 1):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        if s is None:
            p.add_run(f"{i}.  Complete your answers on the document provided and "
                      "upload the completed answers to the LMS at ").font.size = Pt(11)
            add_hyperlink(p, LMS_URL, LMS_URL)
            p.add_run(".").font.size = Pt(11)
        else:
            p.add_run(f"{i}.  {s}").font.size = Pt(11)

def grading(doc, what):
    heading(doc, "Grading")
    para(doc, what, size=11, after=12)
    for ln in ["Grade: _______  (C / NYC)",
               "Assessor Name: __________________________   Assessor NRIC: ________________",
               "Date: ________________________                    Signature: ____________________"]:
        p = para(doc, ln, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

def finish(doc, path):
    prodoc.add_page_numbers(doc); prodoc.enable_update_fields(doc)
    doc.save(path); print("  saved:", os.path.basename(path))

# ---------------------------------------------------------------- builders
def build_wa(answers):
    doc = base_doc()
    kind = "Written Assessment (SAQ) — Answer Key" if answers else "Written Assessment (SAQ)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Answers to Written Assessment (SAQ)" if answers else "Written Assessment (SAQ)",
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        # Page 2 — candidate information, instructions and grading; questions begin on the next page.
        candidate_block(doc); instructions(doc, "1 hour")
        grading(doc, "Candidate has answered all written questions and demonstrated the underpinning "
                     "knowledge required for the course learning outcomes.")
        page_break(doc)
    para(doc, "Short-Answer Questions (Knowledge)", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Answer all questions in your own words. Each question tests underpinning knowledge covered in the "
              "course slides.", size=10.5, italic=True, color=GREY, after=8)
    for i, (crit, ctx, q, pts) in enumerate(WRITTEN, 1):
        para(doc, f"Question {i}:", size=11.5, bold=True, after=2, before=6)
        para(doc, ctx, size=11, after=3)
        para(doc, f"{q}  ({crit})", size=11, bold=True, after=4)
        answer_box(doc, lines=pts if answers else None)
    suffix = A_VER if answers else Q_VER
    name = (f"Answer to WA (SAQ) - {TITLE} - {suffix}.docx" if answers
            else f"WA (SAQ) - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))

def build_pp(answers):
    doc = base_doc()
    kind = "Practical Performance (PP) — Answer Key" if answers else "Practical Performance (PP)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Answers to Practical Performance Assessment" if answers else "Practical Performance Assessment",
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        # Page 2 — candidate information, instructions and grading; the problem begins on the next page.
        candidate_block(doc); instructions(doc, "90 minutes")
        grading(doc, "Candidate has successfully completed all PP tasks and can explain the overall "
                     "functions and features used to achieve them.")
        page_break(doc)
    para(doc, "Practical Problem", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Scenario", size=11.5, bold=True, after=2)
    para(doc, SCENARIO, size=11, after=8)
    for label, crit, prompt, cap, pts in PRACTICAL:
        para(doc, f"{label} ({crit}):", size=11.5, bold=True, after=2, before=6)
        para(doc, prompt, size=11, after=3)
        para(doc, cap, size=10.5, italic=True, color=GREY, after=4)
        answer_box(doc, code=pts if answers else None, height_pt=150)
    suffix = A_VER if answers else Q_VER
    name = (f"Answer to PP Assessment - {TITLE} - {suffix}.docx" if answers
            else f"PP Assessment - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))

if __name__ == "__main__":
    print("Building WSQ assessment set…")
    build_wa(answers=False); build_wa(answers=True)
    build_pp(answers=False); build_pp(answers=True)
    print(f"Done. WA: {len(WRITTEN)} questions · PP: {len(PRACTICAL)} tasks.")
